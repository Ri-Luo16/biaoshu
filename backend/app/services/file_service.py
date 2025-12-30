"""文件处理服务"""
import asyncio
import base64
import gc
import io
import re
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Tuple
from contextlib import suppress

import aiofiles
import aiohttp
import docx
import fitz  # PyMuPDF
import pdfplumber
from docx2python import docx2python
from fastapi import UploadFile

from ..config import settings


class FileService:
    """文件处理服务"""

    # 图片上传配置
    IMAGE_UPLOAD_URL = "https://mt.agnet.top/image/upload"
    IMAGE_UPLOAD_TIMEOUT = 30  # 超时时间（秒）

    @staticmethod
    async def upload_image_to_server(image_data: bytes, filename: str) -> Optional[str]:
        """上传图片到外部服务器"""
        try:
            form_data = aiohttp.FormData()
            form_data.add_field('file',
                              io.BytesIO(image_data),
                              filename=filename,
                              content_type='image/jpeg')

            timeout = aiohttp.ClientTimeout(total=FileService.IMAGE_UPLOAD_TIMEOUT)
            async with aiohttp.ClientSession(timeout=timeout) as session:
                async with session.post(FileService.IMAGE_UPLOAD_URL, data=form_data) as response:
                    if response.status == 200:
                        result = await response.json()
                        return result.get('file_url')
                    print(f"图片上传失败，状态码: {response.status}")
                    return None
        except Exception as e:
            print(f"图片上传异常: {e}")
            return None

    @staticmethod
    def extract_images_from_pdf(file_path: str | Path) -> List[Tuple[bytes, str, int, int]]:
        """从PDF提取图片，返回 (图片数据, 扩展名, 页码, 图片索引) 列表"""
        images = []
        try:
            with fitz.open(str(file_path)) as doc:
                for page_num in range(doc.page_count):
                    page = doc[page_num]
                    for img_index, img in enumerate(page.get_images(full=True)):
                        with suppress(Exception):
                            xref = img[0]
                            with fitz.Pixmap(doc, xref) as pix:
                                # 转换为RGB格式（处理 Alpha 通道或 CMYK）
                                if pix.n - pix.alpha < 4:
                                    if pix.alpha:
                                        with fitz.Pixmap(fitz.csRGB, pix) as pix_rgb:
                                            img_data = pix_rgb.tobytes("jpeg")
                                    else:
                                        img_data = pix.tobytes("jpeg")
                                else:
                                    with fitz.Pixmap(fitz.csRGB, pix) as pix_rgb:
                                        img_data = pix_rgb.tobytes("jpeg")
                                
                                images.append((img_data, "jpg", page_num + 1, img_index + 1))
            return images
        except Exception as e:
            print(f"PDF图片提取失败: {e}")
            return []

    @staticmethod
    def extract_images_from_docx(file_path: str | Path) -> List[Tuple[bytes, str, int]]:
        """从Word文档提取图片，返回 (图片数据, 扩展名, 图片索引) 列表"""
        images = []
        try:
            doc = docx.Document(str(file_path))
            rels = doc.part.rels
            img_index = 0

            for rel in rels.values():
                if "image" in rel.target_ref:
                    with suppress(Exception):
                        img_data = rel.target_part.blob
                        content_type = rel.target_part.content_type
                        ext = 'jpg'
                        if 'png' in content_type: ext = 'png'
                        elif 'gif' in content_type: ext = 'gif'
                        elif 'bmp' in content_type: ext = 'bmp'

                        img_index += 1
                        images.append((img_data, ext, img_index))
            return images
        except Exception as e:
            print(f"Word文档图片提取失败: {e}")
            return []

    @staticmethod
    def _safe_file_cleanup(file_path: str | Path, max_retries: int = 3) -> bool:
        """安全删除文件，带重试机制"""
        path = Path(file_path)
        for attempt in range(max_retries):
            try:
                if path.exists():
                    gc.collect()
                    path.unlink()
                return True
            except Exception as e:
                if attempt == max_retries - 1:
                    print(f"无法删除文件 {file_path}: {e}")
                    return False
                # asyncio.run 在非异步上下文中使用是安全的
                asyncio.run(asyncio.sleep(0.5))
        return True
    
    @staticmethod
    async def save_uploaded_file(file: UploadFile) -> Path:
        """保存上传的文件并返回文件路径（支持MD5去重）"""
        import hashlib
        
        upload_dir = Path(settings.upload_dir)
        upload_dir.mkdir(parents=True, exist_ok=True)

        # 读取文件内容计算 MD5
        content = await file.read()
        file_md5 = hashlib.md5(content).hexdigest()
        
        # 恢复文件指针供后续读取（如果需要）
        # 但这里我们直接用 read 出来的 content 写入，所以不需要 seek
        # 注意：如果 content 很大，这种方式可能会占内存。但在 FastAPI 中 UploadFile 已经是 SpooledTemporaryFile
        
        filename = Path(file.filename or "unknown_file")
        # 使用 MD5 作为文件名的一部分，或者建立映射
        # 这里为了简单直观，保留原名但检查是否存在相同 MD5 的文件
        
        # 简单策略：文件名格式变为 {md5}_{filename}
        # 这样如果同一个文件上传，文件名是一样的
        safe_filename = f"{file_md5}_{filename.name}"
        file_path = upload_dir / safe_filename
        
        if file_path.exists():
            print(f"文件已存在 (MD5命中): {file_path}")
            return file_path

        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(content)

        return file_path
    
    @staticmethod
    async def extract_text_from_pdf(file_path: str | Path) -> str:
        """从PDF文件提取文本"""
        try:
            text = await FileService._extract_pdf_with_pymupdf(file_path)
            if len(text.strip()) < 200:
                plumber_text = await FileService._extract_pdf_with_pdfplumber(file_path)
                if len(plumber_text.strip()) > len(text.strip()):
                    text = plumber_text
            return text
        except Exception as e:
            print(f"高级 PDF 提取失败，尝试基础提取: {e}")
            with suppress(Exception):
                return await FileService._extract_pdf_with_pdfplumber(file_path)
            return FileService._extract_pdf_with_pypdf2(file_path)
    
    @staticmethod
    async def _extract_pdf_with_pdfplumber(file_path: str | Path) -> str:
        """使用pdfplumber提取PDF文本"""
        try:
            extracted_text = []
            image_references = []
            global_img_counter = 1

            all_images = FileService.extract_images_from_pdf(file_path)
            page_images_map = {}
            for img_data, ext, page_num, img_index in all_images:
                page_images_map.setdefault(page_num, []).append((img_data, ext, img_index))

            with pdfplumber.open(str(file_path)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    extracted_text.append(f"\n--- 第 {page_num} 页 ---\n")
                    if text := page.extract_text():
                        img_pattern = r'----.*?(?:image|img|media).*?----'
                        img_matches = list(re.finditer(img_pattern, text, re.IGNORECASE))

                        if img_matches and (page_images := page_images_map.get(page_num)):
                            processed_text = text
                            for match in img_matches:
                                if global_img_counter - 1 < len(page_images):
                                    img_data, ext, img_index = page_images[global_img_counter - 1]
                                    img_name = f"pdf_p{page_num}_i{img_index}.{ext}"
                                    if image_url := await FileService.upload_image_to_server(img_data, img_name):
                                        processed_text = processed_text.replace(match.group(), f"[图片{global_img_counter}]", 1)
                                        image_references.append(f"[图片{global_img_counter}]: {image_url}")
                                        global_img_counter += 1
                            extracted_text.append(processed_text)
                        else:
                            extracted_text.append(text)

                    if tables := page.extract_tables():
                        for table_num, table in enumerate(tables, 1):
                            extracted_text.append(f"\n[表格 {table_num}]")
                            for row in table:
                                if row:
                                    extracted_text.append(" | ".join(str(c) if c else "" for c in row))
                            extracted_text.append("[表格结束]\n")

            if image_references:
                extracted_text.extend(["\n\n--- 图片引用 ---", *image_references])

            result = "\n".join(extracted_text).strip()
            
            if len(result.replace("--- 第", "").strip()) < 100 and all_images:
                fitz_text = await FileService._extract_pdf_with_pymupdf(file_path)
                if len(fitz_text) > len(result):
                    return fitz_text

            return result
        except Exception as e:
            with suppress(Exception):
                return await FileService._extract_pdf_with_pymupdf(file_path)
            raise Exception(f"PDF文件读取失败: {e}") from e
    
    @staticmethod
    async def _extract_pdf_with_pymupdf(file_path: str | Path) -> str:
        """使用PyMuPDF提取PDF文本和图片"""
        try:
            extracted_text = []
            with fitz.open(str(file_path)) as doc:
                for page_num in range(doc.page_count):
                    page = doc[page_num]
                    extracted_text.append(f"\n--- 第 {page_num + 1} 页 ---\n")
                    if text := page.get_text("text", sort=True):
                        extracted_text.append(text)
                    
                    with suppress(Exception):
                        if tabs := page.find_tables():
                            for table_num, table in enumerate(tabs.tables, 1):
                                extracted_text.append(f"\n[表格 {table_num}]")
                                for row in table.extract():
                                    if row:
                                        extracted_text.append(" | ".join(str(c) if c else "" for c in row))
                                extracted_text.append("[表格结束]\n")
            return "\n".join(extracted_text).strip()
        except Exception as e:
            raise Exception(f"PyMuPDF 提取失败: {e}") from e
    
    @staticmethod 
    def _extract_pdf_with_pypdf2(file_path: str | Path) -> str:
        """使用PyPDF2提取PDF文本（原方法）"""
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                return "\n".join(p.extract_text() for p in reader.pages).strip()
        except Exception as e:
            raise Exception(f"PDF文件读取失败: {e}") from e
    
    @staticmethod
    async def extract_text_from_docx(file_path: str | Path) -> str:
        """从Word文档提取文本 (增强版)"""
        try:
            # 1. 尝试使用 docx2python
            return await FileService._extract_docx_with_docx2python(file_path)
        except Exception as e1:
            print(f"docx2python 提取失败: {e1}")
            try:
                # 2. 尝试使用 python-docx
                return await FileService._extract_docx_with_python_docx(file_path)
            except Exception as e2:
                print(f"python-docx 提取失败: {e2}")
                # 3. 尝试使用 win32com (仅限 Windows，处理 .doc 或 伪装 .docx)
                try:
                    return await FileService._extract_word_with_win32com(file_path)
                except Exception as e3:
                    print(f"win32com 提取失败: {e3}")
                    raise Exception(f"Word 文档解析全面失败。请检查文件是否损坏或加密。")

    @staticmethod
    async def _extract_word_with_win32com(file_path: str | Path) -> str:
        """使用 win32com 调用 Word 应用程序提取文本 (Windows Only)"""
        import win32com.client
        import pythoncom
        from pathlib import Path
        
        # 初始化 COM 库 (在异步线程中是必须的)
        pythoncom.CoInitialize()
        
        word = None
        doc = None
        abs_path = str(Path(file_path).resolve())
        
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            
            doc = word.Documents.Open(abs_path, ReadOnly=True)
            text = doc.Range().Text
            
            # 简单的清理
            text = text.replace('\r', '\n').strip()
            return text
            
        except Exception as e:
            raise e
        finally:
            if doc:
                try: doc.Close(0) # 0=wdDoNotSaveChanges
                except: pass
            if word:
                try: word.Quit()
                except: pass
            pythoncom.CoUninitialize()
    
    @staticmethod
    async def _extract_docx_with_docx2python(file_path: str | Path) -> str:
        """使用docx2python提取Word文档内容"""
        try:
            extracted_text = []
            image_references = []
            global_img_counter = 1
            all_images = FileService.extract_images_from_docx(file_path)

            with docx2python(str(file_path)) as content:
                if hasattr(content, 'document'):
                    for section in content.document:
                        for element in section:
                            if isinstance(element, list):
                                extracted_text.append("\n[表格内容]")
                                for row in element:
                                    if isinstance(row, list):
                                        row_text = " | ".join(str(c).strip() for c in row if c)
                                        if row_text: extracted_text.append(row_text)
                                    else:
                                        extracted_text.append(str(row))
                                extracted_text.append("[表格结束]\n")
                            else:
                                text = str(element).strip()
                                if text:
                                    img_matches = list(re.finditer(r'----.*?(?:image|img|media).*?----', text, re.IGNORECASE))
                                    if img_matches and all_images:
                                        processed_text = text
                                        for match in img_matches:
                                            if global_img_counter <= len(all_images):
                                                img_data, ext, _ = all_images[global_img_counter - 1]
                                                if url := await FileService.upload_image_to_server(img_data, f"docx_i{global_img_counter}.{ext}"):
                                                    processed_text = processed_text.replace(match.group(), f"[图片{global_img_counter}]", 1)
                                                    image_references.append(f"[图片{global_img_counter}]: {url}")
                                                    global_img_counter += 1
                                        extracted_text.append(processed_text)
                                    else:
                                        extracted_text.append(text)

            if image_references:
                extracted_text.extend(["\n\n--- 图片引用 ---", *image_references])

            return "\n".join(extracted_text).strip()
        except Exception:
            return await FileService._extract_docx_with_python_docx(file_path)
    
    @staticmethod
    async def _extract_docx_with_python_docx(file_path: str | Path) -> str:
        """使用python-docx提取Word文档内容"""
        try:
            doc = docx.Document(str(file_path))
            extracted_text = []
            image_references = []
            global_img_counter = 1
            all_images = FileService.extract_images_from_docx(file_path)

            for paragraph in doc.paragraphs:
                if text := paragraph.text.strip():
                    img_matches = list(re.finditer(r'----.*?(?:image|img|media).*?----', text, re.IGNORECASE))
                    if img_matches and all_images:
                        processed_text = text
                        for match in img_matches:
                            if global_img_counter <= len(all_images):
                                img_data, ext, _ = all_images[global_img_counter - 1]
                                if url := await FileService.upload_image_to_server(img_data, f"docx_i{global_img_counter}.{ext}"):
                                    processed_text = processed_text.replace(match.group(), f"[图片{global_img_counter}]", 1)
                                    image_references.append(f"[图片{global_img_counter}]: {url}")
                                    global_img_counter += 1
                        extracted_text.append(processed_text)
                    else:
                        extracted_text.append(text)

            for table_num, table in enumerate(doc.tables, 1):
                extracted_text.append(f"\n[表格 {table_num}]")
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text: extracted_text.append(row_text)
                extracted_text.append("[表格结束]\n")

            if image_references:
                extracted_text.extend(["\n\n--- 图片引用 ---", *image_references])

            return "\n".join(extracted_text).strip()
        except Exception as e:
            raise Exception(f"Word文档读取失败: {e}") from e
    
    @staticmethod
    async def process_vectorization_background(text: str, file_path: Path):
        """后台异步处理向量化任务"""
        try:
            from .milvus_service import MilvusService
            from langchain_text_splitters import RecursiveCharacterTextSplitter
            
            print(f"后台任务启动: 开始为 {file_path.name} 进行向量化...")
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                separators=["\n\n", "\n", "。", "！", "？", " ", ""]
            )
            chunks = text_splitter.split_text(text)
            
            if not chunks:
                print("文档内容过少，跳过向量化")
                return

            milvus_service = MilvusService()
            await milvus_service.add_documents(
                texts=chunks,
                metadatas=[{"source": str(file_path.name), "path": str(file_path)}] * len(chunks)
            )
            print(f"后台任务完成: {len(chunks)} 个片段已存入 Milvus")
        except Exception as e:
            print(f"后台向量化任务失败: {e}")

    @staticmethod
    async def process_uploaded_file(file: UploadFile, background_tasks = None) -> Tuple[str, str]:
        """处理上传的文件并提取文本内容，返回 (文本内容, 文件URL)"""
        print(f"开始处理文件: {file.filename}, 类型: {file.content_type}")
        
        # 1. 计算 MD5 并保存文件 (去重)
        import hashlib
        content = await file.read()
        if len(content) > settings.max_file_size:
            raise Exception(f"文件大小超过限制 ({settings.max_file_size / 1024 / 1024}MB)")
        
        file_md5 = hashlib.md5(content).hexdigest()
        filename = Path(file.filename or "unknown_file")
        safe_filename = f"{file_md5}_{filename.name}"
        
        upload_dir = Path(settings.upload_dir)
        upload_dir.mkdir(parents=True, exist_ok=True)
        file_path = upload_dir / safe_filename
        
        # 检查是否已存在
        is_existing_file = False
        if file_path.exists():
            print(f"文件已存在 (MD5命中): {file_path}")
            is_existing_file = True
        else:
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(content)
            print(f"文件已保存至: {file_path}")
            
        file_url = f"/api/uploads/{safe_filename}"
        
        # 2. 检查是否有缓存的解析结果
        # 我们约定：解析后的文本保存在 {filename}.txt 中
        cache_path = file_path.with_suffix(file_path.suffix + ".txt")
        
        if is_existing_file and cache_path.exists():
            print(f"发现缓存的解析结果，直接使用: {cache_path}")
            async with aiofiles.open(cache_path, 'r', encoding='utf-8') as f:
                text = await f.read()
            # 如果缓存存在，说明之前肯定也做过向量化了（或者正在做），这里可以跳过
            print(f"缓存命中，跳过提取和向量化。字数: {len(text)}")
            return text, file_url

        try:
            filename_lower = (file.filename or "").lower()
            is_pdf = file.content_type == "application/pdf" or filename_lower.endswith(".pdf")
            is_docx = (file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or 
                       filename_lower.endswith((".docx", ".docm")))
            is_doc = file.content_type == "application/msword" or filename_lower.endswith(".doc")
            is_image = (file.content_type and file.content_type.startswith("image/")) or filename_lower.endswith(('.png', '.jpg', '.jpeg', '.bmp', '.webp'))
            
            needs_new_file = False
            text = ""
            
            if is_pdf:
                print("检测到 PDF 文件，开始提取...")
                # 1. 尝试快速提取文本层
                text = await FileService.extract_text_from_pdf(file_path)
                
                # 2. 智能判定是否需要 OCR
                clean_text = re.sub(r'\s+', '', text)
                
                if len(clean_text) > 100:
                    print(f"PDF 文本层提取成功，有效字数: {len(clean_text)}，跳过 OCR。")
                else:
                    print(f"PDF 文本层提取内容过少 ({len(clean_text)} 字)，判定为扫描件或纯图片，启动 OCR...")
                    ocr_text = await FileService.perform_ocr_on_pdf(file_path)
                    if len(ocr_text) > len(text):
                        text = ocr_text
                        needs_new_file = True
            
            elif is_docx:
                print("检测到 Word (Docx) 文件，开始提取...")
                text = await FileService.extract_text_from_docx(file_path)
                if not text.strip():
                    if images := FileService.extract_images_from_docx(file_path):
                        print("检测到 Word 文档包含图片但无文字，尝试对图片进行 OCR...")
                        ocr_texts = []
                        from .openai_service import OpenAIService
                        openai_service = OpenAIService()
                        for i, (img_data, _, _) in enumerate(images, 1):
                            base64_img = base64.b64encode(img_data).decode('utf-8')
                            if img_text := await openai_service.ocr_image(base64_img):
                                ocr_texts.append(f"--- 图片 {i} (OCR) ---\n{img_text}")
                        
                        if ocr_texts:
                            text = "\n\n".join(ocr_texts)
                            needs_new_file = True
                        else:
                            raise Exception("该 Word 文档仅包含图片，且 OCR 识别未能提取到文字。")
                    else:
                        raise Exception("未能从该 Word 文档中提取到任何文字内容。")
            
            elif is_doc:
                raise Exception("暂不支持旧版 Word (.doc) 格式。请将其转换为 .docx 格式后再上传。")
            
            elif is_image:
                print(f"检测到图片文件: {file.filename}，开始 OCR 识别...")
                if text := await FileService.perform_ocr_on_image(file_path):
                    needs_new_file = True
                else:
                    raise Exception("无法从该图片中提取文字内容，OCR 识别结果为空。")
            
            else:
                raise Exception(f"不支持的文件类型: {file.content_type}")

            if not text or not text.strip():
                raise Exception("无法从该文件中提取文字内容。")

            if needs_new_file:
                new_pdf_name = f"ocr_{file_path.name}.pdf"
                new_pdf_path = Path(settings.upload_dir) / new_pdf_name
                await FileService.generate_pdf_from_text(text, new_pdf_path)
                file_url = f"/api/uploads/{new_pdf_name}"
                print(f"已生成 OCR 结果 PDF: {new_pdf_path}")

            print(f"文件提取成功，字数: {len(text)}")
            
            # 3. 保存解析结果到缓存文件
            try:
                async with aiofiles.open(cache_path, 'w', encoding='utf-8') as f:
                    await f.write(text)
                print(f"解析结果已缓存至: {cache_path}")
            except Exception as e:
                print(f"缓存写入失败: {e}")

            # 将耗时的向量化操作移入后台任务
            if background_tasks:
                background_tasks.add_task(FileService.process_vectorization_background, text, file_path)
                print("已将向量化任务加入后台队列")
            
            return text, file_url

        except Exception as e:
            # 注意：如果不删除文件，下次上传相同文件可能会因为找不到缓存而再次失败，
            # 但如果文件本身有问题，删除它是对的。
            # 如果是 existing_file，我们不应该删除它，因为它可能在其他地方被引用
            if 'file_path' in locals() and not is_existing_file:
                FileService._safe_file_cleanup(file_path)
            raise e

    @staticmethod
    async def generate_pdf_from_text(text: str, output_path: str | Path) -> None:
        """将文字内容生成 PDF 文件"""
        try:
            with fitz.open() as doc:
                page = doc.new_page()
                font_name = "china-s"
                lines = text.split("\n")
                y_offset, margin, line_height, font_size = 50, 50, 15, 10
                chars_per_line = int((page.rect.width - 2 * margin) / (font_size * 0.8))
                
                for line in lines:
                    if not line.strip():
                        y_offset += line_height
                        continue
                    while len(line) > 0:
                        if y_offset > page.rect.height - margin:
                            page = doc.new_page()
                            y_offset = 50
                        chunk = line[:chars_per_line]
                        line = line[chars_per_line:]
                        with suppress(Exception):
                            page.insert_text((margin, y_offset), chunk, fontname=font_name, fontsize=font_size)
                        y_offset += line_height
                doc.save(str(output_path))
        except Exception as e:
            print(f"生成 PDF 失败: {e}")

    @staticmethod
    async def perform_ocr_on_pdf(file_path: str | Path) -> str:
        """对 PDF 文件进行 OCR 识别 (并发优化版)"""
        from .openai_service import OpenAIService
        openai_service = OpenAIService()
        
        try:
            with fitz.open(str(file_path)) as doc:
                max_pages = 35  
                num_pages = min(doc.page_count, max_pages)
                print(f"开始对 PDF 进行 OCR 识别，共 {num_pages} 页 (并发处理)...")
                
                tasks = []
                for page_num in range(num_pages):
                    page = doc[page_num]
                    # 降低分辨率以加快传输和处理，matrix=1.5 通常足够识别文字
                    pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
                    base64_image = base64.b64encode(pix.tobytes("jpeg")).decode('utf-8')
                    tasks.append(openai_service.ocr_image(base64_image))

            # 并发执行 OCR 请求
            # 注意：这可能会瞬间消耗大量 Token 或触发 API 速率限制
            # 建议分批执行，例如每批 5 页
            batch_size = 5
            results = []
            for i in range(0, len(tasks), batch_size):
                batch = tasks[i:i + batch_size]
                print(f"正在处理 OCR 批次 {i//batch_size + 1}/{(len(tasks)+batch_size-1)//batch_size}...")
                batch_results = await asyncio.gather(*batch, return_exceptions=True)
                results.extend(batch_results)

            full_text = []
            for i, res in enumerate(results):
                if isinstance(res, Exception):
                    print(f"PDF 第 {i + 1} 页 OCR 失败: {res}")
                    full_text.append(f"--- 第 {i + 1} 页 (OCR 失败) ---")
                elif res:
                    full_text.append(f"--- 第 {i + 1} 页 (OCR) ---\n{res}")
                else:
                    full_text.append(f"--- 第 {i + 1} 页 (无内容) ---")
            
            return "\n\n".join(full_text)
        except Exception as e:
            print(f"PDF OCR 异常: {e}")
            return ""

    @staticmethod
    async def perform_ocr_on_image(file_path: str | Path) -> str:
        """对图片文件进行 OCR 识别"""
        from .openai_service import OpenAIService
        openai_service = OpenAIService()
        try:
            path = Path(file_path)
            img_bytes = path.read_bytes()
            base64_image = base64.b64encode(img_bytes).decode('utf-8')
            return await openai_service.ocr_image(base64_image)
        except Exception as e:
            print(f"图片 OCR 失败: {e}")
            return ""
