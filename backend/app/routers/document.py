import json
import io
import re
import traceback
from pathlib import Path
from urllib.parse import quote

import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse

from ..models.schemas import FileUploadResponse, AnalysisRequest, AnalysisType, WordExportRequest
from ..services.file_service import FileService
from ..services.openai_service import OpenAIService
from ..utils.config_manager import config_manager
from ..utils.sse import sse_response

router = APIRouter(prefix="/api/document", tags=["文档处理"])


def set_font(run: docx.text.run.Run, size_pt: float = None) -> None:
    run.font.name = "宋体"
    if size_pt:
        run.font.size = Pt(size_pt)
    r = run._element.rPr
    if r is not None and r.rFonts is not None:
        r.rFonts.set(qn("w:eastAsia"), "宋体")


def set_paragraph_font(paragraph: docx.text.paragraph.Paragraph, size_pt: float = None) -> None:
    for run in paragraph.runs:
        set_font(run, size_pt)


def add_page_number(run: docx.text.run.Run) -> None:
    elements = [
        ('w:fldChar', {'w:fldCharType': 'begin'}),
        ('w:instrText', {'xml:space': 'preserve'}, 'PAGE'),
        ('w:fldChar', {'w:fldCharType': 'separate'}),
        ('w:t', {}, '1'),
        ('w:fldChar', {'w:fldCharType': 'end'}),
    ]
    
    for elem_data in elements:
        elem = OxmlElement(elem_data[0])
        if len(elem_data) > 1:
            for key, value in elem_data[1].items():
                elem.set(qn(key), value)
        if len(elem_data) > 2:
            elem.text = elem_data[2]
        run._element.append(elem)


def add_markdown_runs(para: docx.text.paragraph.Paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        run = para.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.text, run.bold = part[2:-2], True
        elif part.startswith("*") and part.endswith("*"):
            run.text, run.italic = part[1:-1], True
        else:
            run.text = part
        set_font(run)


def add_markdown_paragraph(doc: docx.Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Inches(0.3)
    add_markdown_runs(para, text)


def render_markdown_blocks(doc: docx.Document, blocks: list) -> None:
    for block in blocks:
        kind = block[0]
        if kind == "list":
            for list_kind, num, text in block[1]:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                prefix = "• " if list_kind == "unordered" else f"{num}. "
                run = p.add_run(prefix)
                set_font(run)
                add_markdown_runs(p, text)
        elif kind == "heading":
            h = doc.add_heading(block[2], level=block[1])
            set_paragraph_font(h)
        elif kind == "paragraph":
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.3)
            add_markdown_runs(p, block[1])


def parse_markdown_blocks(content: str) -> list:
    blocks, lines, i = [], content.split("\n"), 0
    while i < len(lines):
        line = lines[i].rstrip("\r").strip()
        if not line: i += 1; continue
        if line.startswith(("- ", "* ")) or re.match(r"^\d+\.\s", line):
            items = []
            while i < len(lines):
                stripped = lines[i].strip()
                if stripped.startswith(("- ", "* ")):
                    items.append(("unordered", None, re.sub(r"^[-*]\s+", "", stripped))); i += 1
                elif m := re.match(r"^(\d+)\.\s+(.*)$", stripped):
                    items.append(("ordered", m.group(1), m.group(2))); i += 1
                else: break
            blocks.append(("list", items))
        elif line.startswith("#"):
            if m := re.match(r"^(#+)\s*(.*)$", line): blocks.append(("heading", min(len(m.group(1)), 3), m.group(2)))
            i += 1
        else:
            p_lines = []
            while i < len(lines):
                stripped = lines[i].strip()
                if stripped and not stripped.startswith(("-", "*", "#")): p_lines.append(stripped); i += 1
                else: break
            if p_lines: blocks.append(("paragraph", " ".join(p_lines)))
    return blocks


@router.post("/upload", response_model=FileUploadResponse)
async def upload_file(file: UploadFile = File(...)) -> FileUploadResponse:
    try:
        allowed_exts = {".pdf", ".docx", ".doc", ".docm"}
        allowed_types = {
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
            "application/vnd.ms-word.document.macroEnabled.12",
            "application/octet-stream"
        }
        
        filename = file.filename or ""
        ext = Path(filename).suffix.lower()
        
        if file.content_type not in allowed_types and ext not in allowed_exts:
            return FileUploadResponse(
                success=False,
                message="不支持的文件类型，请上传 PDF 或 Word (.docx) 文档"
            )
        
        file_content, file_url = await FileService.process_uploaded_file(file)
        
        return FileUploadResponse(
            success=True,
            message=f"文件 {filename} 上传成功",
            filename=filename,
            file_content=file_content,
            file_url=file_url
        )
        
    except Exception as e:
        traceback.print_exc()
        return FileUploadResponse(
            success=False,
            message=f"文件处理失败: {e}"
        )


@router.post("/analyze-stream")
async def analyze_document_stream(request: AnalysisRequest) -> StreamingResponse:
    """流式分析文档内容"""
    try:
        config = config_manager.load_config()
        if not config.get('api_key'):
            raise HTTPException(status_code=400, detail="请先配置OpenAI API密钥")

        openai_service = OpenAIService()
        
        async def generate():
            if request.analysis_type == AnalysisType.OVERVIEW:
                system_prompt = """你是一个专业的标书撰写专家。请分析用户发来的招标文件，提取并总结项目概述信息。
            
请重点关注以下方面：
1. 项目名称和基本信息
2. 项目背景和目的
3. 项目规模和预算
4. 项目时间安排
5. 项目要实施的具体内容
6. 主要技术特点
7. 其他关键要求

工作要求：
1. 保持提取信息的全面性和准确性，尽量使用原文内容，不要自己编写
2. 只关注与项目实施有关的内容，不提取商务信息
3. 直接返回整理好的项目概述，除此之外不返回任何其他内容
"""
            elif request.analysis_type == AnalysisType.REQUIREMENTS:
                system_prompt = """你是一名专业的招标文件分析师，擅长从复杂的招标文档中高效提取“技术评分项”相关内容。
对每一项技术评分项，按以下结构化格式输出（若信息缺失，标注“未提及”）：
【评分项名称】：<原文描述，保留专业术语>
【权重/分值】：<具体分值或占比，必须取整，不要保留小数位。例如“30分”或“40%”>
【评分标准】：<详细规则，如“≥95%得满分，每低1%扣0.5分”>
【数据来源】：<文档中的位置，如“第5.2.3条”或“附件3-表2”>

直接返回提取结果，除此之外不输出任何其他内容
"""
            else:  # structural
                system_prompt = """你是一名资深招投标分析师，请对用户提供的招标文件进行全维度结构化解析，并输出为详细的 JSON 格式。
解析内容需包含以下核心板块：

1. executive_summary: 用一句话概括本项目的核心本质和对投标方的核心要求（如：“本项目是一个高预算、高技术门槛的政府软件采购项目，侧重于系统安全性和本地化服务能力”）。

2. action_items: 列出 3-5 条投标方必须立即关注的“核心行动项/关键需求”，以短句形式表达。

3. project_summary (项目摘要看板):
   - total_score: 总分值（如 100，必须为整数）
   - mandatory_count: 强制性/星号条款总数
   - budget: 预算金额（若提及）
   - delivery_time: 交付/工期/服务期要求
   - qualification_count: 核心资格要求数量

4. qualification (资格门槛):
   - 提取所有参与本项目必须具备的资质、业绩、人员等硬性要求。

5. technical_requirements (技术需求清单):
   - 包含 item (需求项名称), value (具体参数/要求), mandatory (是否为强制性条款，布尔值)。

6. scoring_criteria (评分权重概览):
   - 提取各板块（如技术方案、商务、价格、售后）的得分占比（分值必须取整，不保留小数）。

7. implicit_needs (隐性需求挖掘):
   - 挖掘文字背后的深层需求，如“偏向本地化服务”、“对兼容性要求极高”等。

8. risk_flags (核心风险预警):
   - 识别可能导致废标或重大执行风险的条款。

请直接返回纯 JSON 格式的内容，不要包含 Markdown 代码块标签（如 ```json）或任何解释性文字。
"""
            
            mapping = {
                AnalysisType.OVERVIEW: "项目概述",
                AnalysisType.REQUIREMENTS: "技术评分要求",
                AnalysisType.STRUCTURAL: "结构化解析"
            }
            analysis_type_cn = mapping.get(request.analysis_type, "分析")
            user_prompt = f"请分析以下招标文件内容，提取{analysis_type_cn}信息：\n\n{request.file_content}"
            
            messages = [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}]
            async for chunk in openai_service.stream_chat_completion(messages, temperature=0.3):
                yield f"data: {json.dumps({'chunk': chunk}, ensure_ascii=False)}\n\n"
            yield "data: [DONE]\n\n"
        
        return sse_response(generate())
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文档分析失败: {e}")


@router.post("/export-word")
async def export_word(request: WordExportRequest) -> StreamingResponse:
    """根据目录数据导出Word文档（标准标书整合版）"""
    try:
        doc = docx.Document()
        
        # 1. 基础样式优化
        styles = doc.styles
        for name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Title"]:
            if name in styles:
                style = styles[name]
                style.font.name = "宋体"
                if style._element.rPr is None: style._element._add_rPr()
                style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                
                if name == "Normal":
                    style.font.size = Pt(12)  # 小四
                    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                    style.paragraph_format.space_after = Pt(6)
                elif name == "Heading 1":
                    style.font.size = Pt(16)
                    style.font.bold = True
                    style.paragraph_format.space_before = Pt(12)
                    style.paragraph_format.space_after = Pt(6)
                elif name == "Heading 2":
                    style.font.size = Pt(14)
                    style.font.bold = True
                    style.paragraph_format.space_before = Pt(10)
                    style.paragraph_format.space_after = Pt(4)

        # 2. 标书封面生成
        doc.add_paragraph("\n\n\n")
        title_p = doc.add_paragraph()
        title_run = title_p.add_run("投 标 文 件")
        title_run.bold, title_run.font.size = True, Pt(42)
        set_font(title_run)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("\n")
        tech_p = doc.add_paragraph()
        tech_run = tech_p.add_run("（技 术 部 分）")
        tech_run.font.size = Pt(22)
        set_font(tech_run)
        tech_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("\n\n\n\n")
        
        # 封面信息表格
        info_table = doc.add_table(rows=4, cols=2)
        info_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        labels = ["项 目 名 称：", "项目编号：", "投 标 人：", "日    期："]
        values = [
            request.project_name or "————",
            request.project_number or "————",
            request.bidder_name or "————",
            request.bid_date or "202X年XX月XX日"
        ]
        for i in range(4):
            cell_l = info_table.cell(i, 0)
            cell_l.text = labels[i]
            p_l = cell_l.paragraphs[0]
            p_l.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_paragraph_font(p_l, 14)
            
            cell_v = info_table.cell(i, 1)
            cell_v.text = values[i]
            p_v = cell_v.paragraphs[0]
            set_paragraph_font(p_v, 14)
            for run in p_v.runs: run.underline = True

        doc.add_page_break()

        # 3. 页码设置 (从目录页开始)
        for section in doc.sections:
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = footer_para.add_run()
            set_font(run, 9)
            add_page_number(run)

        # 4. 目录页
        doc.add_heading("目  录", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        def add_toc_item(items, level=1):
            for item in items:
                toc_p = doc.add_paragraph()
                toc_p.paragraph_format.left_indent = Inches(0.2 * (level - 1))
                toc_run = toc_p.add_run(f"{item.id} {item.title}")
                set_font(toc_run, 10.5)
                if level == 1: toc_run.bold = True
                if item.children: add_toc_item(item.children, level + 1)
        
        add_toc_item(request.outline)
        doc.add_page_break()

        # 5. 项目概述
        if request.project_overview:
            h = doc.add_heading("项目概述", level=1)
            set_paragraph_font(h)
            render_markdown_blocks(doc, parse_markdown_blocks(request.project_overview))

        def add_outline_items(items: list, level: int = 1) -> None:
            for item in items:
                if level == 1:
                    doc.add_page_break()
                h = doc.add_heading(f"{item.id} {item.title}", level=min(level, 3))
                set_paragraph_font(h)
                if not item.children:
                    if c := (item.content or "").strip(): 
                        render_markdown_blocks(doc, parse_markdown_blocks(c))
                else:
                    add_outline_items(item.children, level + 1)

        add_outline_items(request.outline)

        # 7. 落款与盖章页
        doc.add_page_break()
        doc.add_paragraph("\n\n\n")
        signature_lines = [
            "投标人（盖章）：__________________________",
            "\n",
            "法定代表人或授权代表（签字）：________________",
            "\n",
            f"日    期：{request.bid_date or '202X年XX月XX日'}"
        ]
        for text in signature_lines:
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_font(run, 14)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        encoded_filename = quote(f"{request.project_name or '标书文档'}.docx")
        return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"})
    except Exception as e:
        print(f"导出Word失败: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"导出Word失败: {e}")
