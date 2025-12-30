import os
from pathlib import Path
from typing import Dict, Any, Union
import docx
from docx.shared import Pt
from app.config import settings

def export_response_to_docx(content_structure: Dict[str, Any], file_name: str = "tender_response.docx") -> str:
    """
    将投标响应内容导出为 Word 文档
    
    Args:
        content_structure: 结构化内容，例如:
            {
                "第一章 投标函": "内容...",
                "第二章 技术方案": {
                    "2.1 项目理解": "内容...",
                    "2.2 总体设计": "内容..."
                }
            }
        file_name: 输出文件名
    
    Returns:
        生成的文件的相对路径 (用于下载)
    """
    try:
        doc = docx.Document()
        
        # 设置基础样式
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        
        # 递归写入内容
        _write_section_recursive(doc, content_structure, level=1)
        
        # 保存文件
        upload_dir = Path(settings.upload_dir)
        upload_dir.mkdir(parents=True, exist_ok=True)
        file_path = upload_dir / file_name
        
        doc.save(file_path)
        return f"/api/uploads/{file_name}"
        
    except Exception as e:
        print(f"Word 导出失败: {e}")
        raise e

def _write_section_recursive(doc, content: Union[str, Dict[str, Any]], level: int):
    """递归写入章节"""
    if isinstance(content, str):
        # 写入文本内容
        for paragraph in content.split('\n'):
            if paragraph.strip():
                p = doc.add_paragraph(paragraph.strip())
                # 简单处理：如果是 # 开头，作为标题
                if paragraph.strip().startswith('# '):
                    p.style = 'Heading 1'
                    p.text = paragraph.strip()[2:]
                elif paragraph.strip().startswith('## '):
                    p.style = 'Heading 2'
                    p.text = paragraph.strip()[3:]
    elif isinstance(content, dict):
        for title, sub_content in content.items():
            # 写入标题
            doc.add_heading(title, level=min(level, 9))
            # 递归写入子内容
            _write_section_recursive(doc, sub_content, level + 1)

